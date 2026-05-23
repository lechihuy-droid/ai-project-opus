"""DOCX template rendering for weekly reports."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
import yaml

from insighthub.schema import Report


SECTIONS: tuple[tuple[str, str], ...] = (
    ("exec_summary", "Executive Summary"),
    ("progress", "Progress"),
    ("completed", "Completed"),
    ("in_progress", "In Progress"),
    ("next_week", "Next Week"),
    ("blockers", "Blockers"),
    ("bugs", "Bugs"),
    ("decisions", "Decisions"),
    ("metrics", "Metrics"),
)
REGISTRY_PATH = Path("templates/registry.yaml")


@dataclass(frozen=True)
class TemplateSpec:
    """Registry entry describing one selectable DOCX template."""

    name: str
    report_type: str
    path: Path
    default: bool = False
    placeholders: tuple[str, ...] = ()


def load_registry(registry_path: Path = REGISTRY_PATH) -> dict[str, TemplateSpec]:
    """Load the template registry from YAML."""

    payload = yaml.safe_load(registry_path.read_text(encoding="utf-8")) or {}
    templates = payload.get("templates", [])
    registry: dict[str, TemplateSpec] = {}
    for item in templates:
        if not isinstance(item, dict):
            continue
        name = str(item.get("name", "")).strip()
        if not name:
            continue
        registry[name] = TemplateSpec(
            name=name,
            report_type=str(item.get("report_type", "weekly")).strip() or "weekly",
            path=Path(str(item.get("path", "")).strip()),
            default=bool(item.get("default", False)),
            placeholders=tuple(str(value) for value in item.get("placeholders", [])),
        )
    if not registry:
        raise ValueError(f"No templates defined in {registry_path}")
    return registry


def resolve_template(report_type: str, template_name: str | None = None) -> TemplateSpec:
    """Return the matching template entry for the requested report type."""

    registry = load_registry()
    if template_name:
        spec = registry.get(template_name)
        if spec is None:
            raise ValueError(f"Unknown template '{template_name}'")
        if spec.report_type != report_type:
            raise ValueError(
                f"Template '{template_name}' does not support report type '{report_type}'"
            )
        return spec

    for spec in registry.values():
        if spec.report_type == report_type and spec.default:
            return spec
    for spec in registry.values():
        if spec.report_type == report_type:
            return spec
    raise ValueError(f"No template registered for report type '{report_type}'")


def render(report: Report, template_path: str) -> DocumentObject:
    """Render a report into a python-docx document."""

    document = Document(template_path)
    section_bodies = {
        section.section_id: section.body or "(no data)"
        for section in report.sections
    }
    replacements = {
        "{{project_name}}": report.project_name,
        "{{period}}": f"{report.period_start.isoformat()} - {report.period_end.isoformat()}",
        "{{overall_status}}": report.overall_status,
    }

    for section_id, _title in SECTIONS:
        replacements[f"{{{{{section_id}}}}}"] = section_bodies.get(section_id, "(no data)")

    for paragraph in list(document.paragraphs):
        text = paragraph.text
        matched = [placeholder for placeholder in replacements if placeholder in text]
        if not matched:
            continue

        if len(matched) == 1 and matched[0].strip("{}") in dict(SECTIONS):
            _replace_with_markdown(paragraph, replacements[matched[0]])
            continue

        rendered = text
        for placeholder in matched:
            rendered = rendered.replace(placeholder, replacements[placeholder])
        _set_paragraph_text(paragraph, rendered)

    template_section_ids = {section_id for section_id, _title in SECTIONS}
    extra_sections = [section for section in report.sections if section.section_id not in template_section_ids]
    for section in extra_sections:
        document.add_heading(section.title, level=1)
        paragraph = document.add_paragraph()
        _replace_with_markdown(paragraph, section.body or "(no data)")

    return document


def _replace_with_markdown(paragraph: Paragraph, markdown: str) -> None:
    lines = [line for line in markdown.splitlines() if line.strip()]
    if not lines:
        lines = ["(no data)"]

    first_text, first_style = _line_to_text_style(lines[0])
    _set_paragraph_text(paragraph, first_text)
    paragraph.style = first_style

    cursor = paragraph
    for line in lines[1:]:
        text, style = _line_to_text_style(line)
        cursor = _insert_paragraph_after(cursor, text, style)


def _line_to_text_style(line: str) -> tuple[str, str]:
    stripped = line.strip()
    if stripped.startswith("- "):
        return stripped[2:].strip(), "List Bullet"
    return stripped, "Normal"


def _set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def _insert_paragraph_after(paragraph: Paragraph, text: str, style: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    new_paragraph.style = style
    new_paragraph.add_run(text)
    return new_paragraph
