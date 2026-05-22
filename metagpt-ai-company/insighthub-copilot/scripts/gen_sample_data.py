"""Generate a synthetic but coherent sample data package for Project Sakura.

The real hackathon ships Sample_Jira_Export.xlsx etc.; until those files are
dropped into data/sample/ this generator produces equivalent files with the
same schema. The data deliberately seeds one instance of every anomaly rule so
the reconciliation engine can be tested for >=85% detection.

Run:  python scripts/gen_sample_data.py
"""

from __future__ import annotations

import json
from datetime import date, datetime, timedelta
from pathlib import Path

import openpyxl
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
SAMPLE = ROOT / "data" / "sample"
MINUTES = SAMPLE / "minutes"

TEAM = ["Tanaka", "Suzuki", "Nguyen", "Pham", "Sato"]
PERIOD_START = date(2026, 5, 15)
PERIOD_END = date(2026, 5, 21)


def d(s: str) -> date:
    return date.fromisoformat(s)


# --------------------------------------------------------------------------
# Jira issues  — SAKURA-1 .. SAKURA-36
# --------------------------------------------------------------------------
def build_issues() -> list[dict]:
    issues: list[dict] = []

    def add(**kw):
        kw.setdefault("priority", "Medium")
        kw.setdefault("labels", "")
        kw.setdefault("reopened", "No")
        kw.setdefault("resolved", None)
        kw.setdefault("sprint", "Sprint 9")
        issues.append(kw)

    # --- 8 normal Done stories/tasks resolved inside the report week --------
    done_specs = [
        ("SAKURA-1", "Implement product search API", "Story", 8, "Tanaka", "2026-05-18"),
        ("SAKURA-2", "Build cart page UI", "Story", 5, "Tanaka", "2026-05-19"),
        ("SAKURA-3", "Checkout API integration", "Story", 8, "Tanaka", "2026-05-20"),
        ("SAKURA-4", "Product detail page", "Task", 3, "Suzuki", "2026-05-17"),
        ("SAKURA-5", "Wishlist feature", "Story", 5, "Tanaka", "2026-05-21"),
        ("SAKURA-6", "Order history screen", "Task", 3, "Nguyen", "2026-05-16"),
        ("SAKURA-7", "Coupon code validation", "Task", 2, "Pham", "2026-05-19"),
        ("SAKURA-8", "Email receipt template", "Task", 2, "Sato", "2026-05-20"),
    ]
    for key, summ, typ, sp, asg, res in done_specs:
        add(key=key, summary=summ, issue_type=typ, status="Done", assignee=asg,
            story_points=sp, created="2026-05-04", updated=res, resolved=res)

    # --- 10 normal In Progress ---------------------------------------------
    ip_specs = [
        ("SAKURA-9", "Payment gateway integration", "Story", 8, "Suzuki"),
        ("SAKURA-10", "Inventory sync service", "Story", 8, "Nguyen"),
        ("SAKURA-11", "Address book management", "Task", 3, "Pham"),
        ("SAKURA-13", "Shipping cost calculator", "Task", 5, "Sato"),
        ("SAKURA-14", "Product recommendation widget", "Story", 5, "Suzuki"),
        ("SAKURA-16", "Admin order dashboard", "Story", 8, "Nguyen"),
        ("SAKURA-17", "Refund processing flow", "Task", 5, "Pham"),
        ("SAKURA-18", "Multi-currency display", "Task", 3, "Sato"),
        ("SAKURA-19", "Guest checkout flow", "Story", 5, "Suzuki"),
        ("SAKURA-21", "Loyalty points engine", "Story", 8, "Nguyen"),
    ]
    for key, summ, typ, sp, asg in ip_specs:
        add(key=key, summary=summ, issue_type=typ, status="In Progress", assignee=asg,
            story_points=sp, created="2026-05-05", updated="2026-05-20")

    # --- 4 normal To Do ----------------------------------------------------
    todo_specs = [
        ("SAKURA-22", "Performance load testing", "Task", 5, "Pham"),
        ("SAKURA-23", "Accessibility audit", "Task", 3, "Sato"),
        ("SAKURA-24", "SEO metadata pass", "Task", 2, "Suzuki"),
        ("SAKURA-34", "Final UAT preparation", "Task", 3, "Nguyen"),
    ]
    for key, summ, typ, sp, asg in todo_specs:
        add(key=key, summary=summ, issue_type=typ, status="To Do", assignee=asg,
            story_points=sp, created="2026-05-10", updated="2026-05-12")

    # --- SEEDED ANOMALIES --------------------------------------------------
    # ANOM-PG-001: Done ticket with no code activity (no commit references it)
    add(key="SAKURA-12", summary="Configure CDN caching rules", issue_type="Task",
        status="Done", assignee="Sato", story_points=3,
        created="2026-05-06", updated="2026-05-19", resolved="2026-05-19")

    # ANOM-PG-002: ticket still In Progress although a merged PR references it
    add(key="SAKURA-15", summary="Order confirmation webhook", issue_type="Story",
        status="In Progress", assignee="Suzuki", story_points=5,
        created="2026-05-05", updated="2026-05-20")

    # ANOM-PG-003: linked to a WBS task already past its planned end date
    add(key="SAKURA-20", summary="Tax calculation module", issue_type="Story",
        status="In Progress", assignee="Pham", story_points=8,
        created="2026-05-04", updated="2026-05-21")

    # ANOM-PG-005: orphan issue — no WBS parent, distinctive summary
    add(key="SAKURA-30", summary="Spike: evaluate GraphQL gateway", issue_type="Task",
        status="In Progress", assignee="Tanaka", story_points=3,
        created="2026-05-12", updated="2026-05-20")

    # ANOM-BG-001: bug aging — open >14 days, no recent change
    add(key="SAKURA-25", summary="Cart total miscalculates with coupon", issue_type="Bug",
        status="In Progress", priority="High", assignee="Pham", story_points=2,
        created="2026-04-20", updated="2026-04-22")

    # ANOM-BG-002: Severity-1 bug unresolved at period end
    add(key="SAKURA-26", summary="Checkout crashes on Safari", issue_type="Bug",
        status="In Progress", priority="Severity-1", assignee="Suzuki", story_points=3,
        created="2026-05-17", updated="2026-05-21")

    # ANOM-BG-003: reopen spike — 3 reopened bugs this period
    for key, summ, asg in [
        ("SAKURA-27", "Search returns stale results", "Nguyen"),
        ("SAKURA-28", "Wishlist item duplicates", "Sato"),
        ("SAKURA-29", "Address validation false negative", "Pham"),
    ]:
        add(key=key, summary=summ, issue_type="Bug", status="In Progress",
            priority="Medium", assignee=asg, story_points=2,
            created="2026-05-08", updated="2026-05-20", reopened="Yes")

    # Normal closed bugs (so reopen rate / bug metrics look realistic)
    for key, summ, asg, res in [
        ("SAKURA-31", "Footer link 404", "Sato", "2026-05-16"),
        ("SAKURA-32", "Image lazy-load flicker", "Suzuki", "2026-05-18"),
        ("SAKURA-33", "Currency symbol misaligned", "Nguyen", "2026-05-20"),
    ]:
        add(key=key, summary=summ, issue_type="Bug", status="Done",
            priority="Low", assignee=asg, story_points=1,
            created="2026-05-11", updated=res, resolved=res)

    # Two more open bugs, normal
    add(key="SAKURA-35", summary="Coupon field accepts whitespace", issue_type="Bug",
        status="In Progress", priority="Low", assignee="Pham", story_points=1,
        created="2026-05-15", updated="2026-05-20")
    add(key="SAKURA-36", summary="Slow first paint on category page", issue_type="Bug",
        status="To Do", priority="Medium", assignee="Sato", story_points=2,
        created="2026-05-18", updated="2026-05-19")

    return issues


# Sprints — ANOM-SC-001 needs two consecutive sprints below 70%.
SPRINTS = [
    ("Sprint 7", "2026-04-17", "2026-04-30", 40, 24),   # 60%
    ("Sprint 8", "2026-05-01", "2026-05-14", 45, 27),   # 60%
    ("Sprint 9", "2026-05-15", "2026-05-28", 48, 31),   # current
]


def write_jira(issues: list[dict]) -> None:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Issues"
    cols = ["Key", "Summary", "Type", "Status", "Priority", "Assignee",
            "Story Points", "Created", "Updated", "Resolved", "Sprint",
            "Labels", "Reopened"]
    ws.append(cols)
    for it in issues:
        ws.append([
            it["key"], it["summary"], it["issue_type"], it["status"],
            it["priority"], it["assignee"], it["story_points"],
            d(it["created"]), d(it["updated"]),
            d(it["resolved"]) if it["resolved"] else None,
            it["sprint"], it["labels"], it["reopened"],
        ])
    sp = wb.create_sheet("Sprints")
    sp.append(["Sprint", "Start", "End", "Committed SP", "Completed SP"])
    for name, s, e, c, comp in SPRINTS:
        sp.append([name, d(s), d(e), c, comp])
    SAMPLE.mkdir(parents=True, exist_ok=True)
    wb.save(SAMPLE / "Sample_Jira_Export.xlsx")


# --------------------------------------------------------------------------
# WBS
# --------------------------------------------------------------------------
def write_wbs() -> None:
    # planned_pct for P2 set high (~85) while Jira actuals are low -> phase drift.
    rows = [
        # task_id, phase, name, p_start, p_end, mm, planned_pct, owner, jira_key
        ("P1-T01", "Phase 1 Design", "Requirements finalization", "2026-04-01", "2026-04-10", 1.5, 100, "Nguyen", ""),
        ("P1-T02", "Phase 1 Design", "UI/UX wireframes", "2026-04-08", "2026-04-20", 2.0, 100, "Sato", ""),
        ("P1-T03", "Phase 1 Design", "Technical architecture", "2026-04-15", "2026-04-30", 1.5, 100, "Tanaka", ""),
        ("P2-T04", "Phase 2 Development", "Product search API", "2026-05-01", "2026-05-18", 2.0, 90, "Tanaka", "SAKURA-1"),
        ("P2-T05", "Phase 2 Development", "Tax calculation module", "2026-05-05", "2026-05-18", 1.5, 85, "Pham", "SAKURA-20"),
        ("P2-T06", "Phase 2 Development", "Payment gateway integration", "2026-05-06", "2026-05-26", 2.5, 80, "Suzuki", "SAKURA-9"),
        ("P2-T07", "Phase 2 Development", "Inventory sync service", "2026-05-08", "2026-05-27", 2.0, 80, "Nguyen", "SAKURA-10"),
        ("P2-T08", "Phase 2 Development", "Cart and checkout", "2026-05-04", "2026-05-22", 2.5, 88, "Tanaka", "SAKURA-3"),
        ("P2-T09", "Phase 2 Development", "Order confirmation webhook", "2026-05-05", "2026-05-25", 1.5, 82, "Suzuki", "SAKURA-15"),
        ("P3-T10", "Phase 3 Testing", "Performance load testing", "2026-05-25", "2026-06-08", 2.0, 0, "Pham", "SAKURA-22"),
        ("P3-T11", "Phase 3 Testing", "Accessibility audit", "2026-05-27", "2026-06-05", 1.0, 0, "Sato", "SAKURA-23"),
        ("P3-T12", "Phase 3 Testing", "User acceptance testing", "2026-06-01", "2026-06-12", 1.5, 0, "Nguyen", "SAKURA-34"),
    ]
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "WBS"
    ws.append(["Task ID", "Phase", "Task Name", "Planned Start", "Planned End",
               "Planned MM", "Planned %", "Owner", "Jira Key"])
    for tid, ph, nm, ps, pe, mm, pct, ow, jk in rows:
        ws.append([tid, ph, nm, d(ps), d(pe), mm, pct, ow, jk])
    wb.save(SAMPLE / "Sample_WBS.xlsx")


# --------------------------------------------------------------------------
# Slack messages
# --------------------------------------------------------------------------
def write_slack() -> None:
    msgs = [
        ("C-001", "Tanaka", "2026-05-15T09:10:00", "Morning all, starting on the search API today. SAKURA-1 on track."),
        ("C-002", "Suzuki", "2026-05-15T10:30:00", "Payment gateway sandbox is slow but progressing on SAKURA-9."),
        ("C-003", "Nguyen", "2026-05-16T11:00:00", "Order history screen SAKURA-6 is done and merged."),
        ("C-004", "Pham", "2026-05-16T14:20:00", "We are blocked on the payment gateway sandbox access from the customer side, cannot proceed with end-to-end testing."),
        ("C-005", "Sato", "2026-05-17T09:45:00", "Product detail page SAKURA-4 finished, moving to CDN config."),
        ("C-006", "Tanaka", "2026-05-18T13:00:00", "Search API SAKURA-1 merged. Demo ready."),
        ("C-007", "Suzuki", "2026-05-18T16:10:00", "Heads up: checkout crashes on Safari, raised SAKURA-26 as Severity-1."),
        ("C-008", "Nguyen", "2026-05-19T10:05:00", "Inventory sync SAKURA-10 needs another day, integration is tricky."),
        ("C-009", "Pham", "2026-05-19T15:30:00", "Coupon validation SAKURA-7 done."),
        ("C-010", "Tanaka", "2026-05-20T09:00:00", "Checkout API SAKURA-3 merged, wishlist SAKURA-5 almost there."),
        ("C-011", "Sato", "2026-05-20T11:40:00", "Email receipt SAKURA-8 completed."),
        ("C-012", "Suzuki", "2026-05-21T14:00:00", "Safari crash SAKURA-26 still under investigation, escalation likely if not fixed Monday."),
        ("C-013", "Nguyen", "2026-05-21T17:00:00", "Wrapping up the week, good progress overall."),
    ]
    data = {"channel": "#project-sakura", "messages": [
        {"id": cid, "channel": "#project-sakura", "user": u, "ts": ts, "text": tx}
        for cid, u, ts, tx in msgs
    ]}
    (SAMPLE / "Sample_Slack_Messages.json").write_text(
        json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")


# --------------------------------------------------------------------------
# GitHub activity
# --------------------------------------------------------------------------
def write_github() -> None:
    # Commits — every normal Done ticket gets one; SAKURA-12 deliberately gets none.
    commit_specs = [
        ("a1b2c01", "SAKURA-1: implement product search API", "Tanaka", "2026-05-18T08:00:00"),
        ("a1b2c02", "SAKURA-2: cart page UI", "Tanaka", "2026-05-19T08:00:00"),
        ("a1b2c03", "SAKURA-3: checkout API integration", "Tanaka", "2026-05-20T08:00:00"),
        ("a1b2c04", "SAKURA-4: product detail page", "Suzuki", "2026-05-17T08:00:00"),
        ("a1b2c05", "SAKURA-5: wishlist feature", "Tanaka", "2026-05-21T08:00:00"),
        ("a1b2c06", "SAKURA-6: order history screen", "Nguyen", "2026-05-16T08:00:00"),
        ("a1b2c07", "SAKURA-7: coupon code validation", "Pham", "2026-05-19T08:00:00"),
        ("a1b2c08", "SAKURA-8: email receipt template", "Sato", "2026-05-20T08:00:00"),
        ("a1b2c09", "SAKURA-15: order confirmation webhook handler", "Suzuki", "2026-05-19T08:00:00"),
        ("a1b2c10", "SAKURA-31: fix footer 404", "Sato", "2026-05-16T08:00:00"),
        ("a1b2c11", "SAKURA-32: fix lazy-load flicker", "Suzuki", "2026-05-18T08:00:00"),
        ("a1b2c12", "SAKURA-33: fix currency alignment", "Nguyen", "2026-05-20T08:00:00"),
        ("a1b2c13", "SAKURA-9: WIP payment gateway client", "Suzuki", "2026-05-20T08:00:00"),
    ]
    commits = [
        {"sha": s, "message": m, "author": a, "date": dt}
        for s, m, a, dt in commit_specs
    ]
    # Pull requests
    pr_specs = [
        # num, title, author, reviewers, state, created, merged, ci
        (1, "SAKURA-1: product search API", "Tanaka", ["Suzuki"], "merged", "2026-05-17T09:00:00", "2026-05-18T10:00:00", "success"),
        (2, "SAKURA-3: checkout API", "Tanaka", ["Nguyen"], "merged", "2026-05-19T09:00:00", "2026-05-20T10:00:00", "success"),
        (3, "SAKURA-6: order history", "Nguyen", ["Pham"], "merged", "2026-05-15T09:00:00", "2026-05-16T10:00:00", "success"),
        (4, "SAKURA-4: product detail page", "Suzuki", ["Sato"], "merged", "2026-05-16T09:00:00", "2026-05-17T10:00:00", "success"),
        # ANOM-PG-002: merged PR references SAKURA-15 but that ticket is still In Progress
        (5, "SAKURA-15: order confirmation webhook", "Suzuki", ["Tanaka"], "merged", "2026-05-18T09:00:00", "2026-05-19T10:00:00", "success"),
        (6, "SAKURA-2: cart page UI", "Tanaka", ["Suzuki"], "merged", "2026-05-18T09:00:00", "2026-05-19T11:00:00", "success"),
        # ANOM-QL-001: merged PR with no reviewer
        (7, "SAKURA-5: wishlist feature", "Tanaka", [], "merged", "2026-05-20T09:00:00", "2026-05-21T10:00:00", "success"),
        # ANOM-QL-002: failing CI on the default branch
        (8, "SAKURA-9: payment gateway integration", "Suzuki", ["Nguyen"], "merged", "2026-05-20T09:00:00", "2026-05-21T12:00:00", "failure"),
        (9, "SAKURA-7: coupon validation", "Pham", ["Sato"], "merged", "2026-05-18T09:00:00", "2026-05-19T13:00:00", "success"),
        (10, "SAKURA-8: email receipt", "Sato", ["Pham"], "open", "2026-05-20T09:00:00", None, "success"),
    ]
    prs = [
        {"number": n, "title": t, "author": a, "reviewers": rv, "state": st,
         "created_at": cr, "merged_at": mg, "ci_status": ci}
        for n, t, a, rv, st, cr, mg, ci in pr_specs
    ]
    data = {"repo": "fpt-japan/project-sakura", "default_branch": "main",
            "commits": commits, "pull_requests": prs}
    (SAMPLE / "Sample_GitHub_Activity.json").write_text(
        json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")


# --------------------------------------------------------------------------
# Meeting minutes (DOCX)
# --------------------------------------------------------------------------
def _minutes_doc(path: Path, title: str, date_str: str,
                 decisions: list[str], actions: list[str], questions: list[str]) -> None:
    doc = Document()
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"Date: {date_str}")
    doc.add_paragraph("Project: Project Sakura")
    doc.add_heading("Decisions", level=1)
    for x in decisions:
        doc.add_paragraph(x, style="List Bullet")
    doc.add_heading("Action Items", level=1)
    for x in actions:
        doc.add_paragraph(x, style="List Bullet")
    doc.add_heading("Questions", level=1)
    for x in questions:
        doc.add_paragraph(x, style="List Bullet")
    doc.save(path)


def write_minutes() -> None:
    MINUTES.mkdir(parents=True, exist_ok=True)
    _minutes_doc(
        MINUTES / "minutes_2026-05-15_sprint-planning.docx",
        "Sprint 9 Planning & Customer QA", "2026-05-15",
        decisions=["Confirmed Sprint 9 scope covering checkout, payments and inventory."],
        actions=["Tanaka to set up the staging environment by 2026-05-20."],
        # ANOM-CS-001: customer question left unanswered
        questions=["Customer asked: when can we expect the load-test environment to be ready?"],
    )
    _minutes_doc(
        MINUTES / "minutes_2026-05-16_customer-sync.docx",
        "Customer Weekly Sync", "2026-05-16",
        # ANOM-RK-002: decision with no follow-up Jira ticket or Slack discussion
        decisions=["We agreed to add a guest-checkout flow as a new scope item."],
        actions=["Nguyen to prepare the API specification draft by 2026-05-22."],
        questions=[],
    )


# --------------------------------------------------------------------------
# Ground truth — which rules were seeded
# --------------------------------------------------------------------------
GROUND_TRUTH = {
    "expected_rules": [
        "ANOM-PG-001", "ANOM-PG-002", "ANOM-PG-003", "ANOM-PG-004", "ANOM-PG-005",
        "ANOM-BG-001", "ANOM-BG-002", "ANOM-BG-003",
        "ANOM-RK-001", "ANOM-RK-002", "ANOM-SC-001",
        "ANOM-QL-001", "ANOM-QL-002", "ANOM-RS-001", "ANOM-CS-001",
    ],
    "seeded": {
        "ANOM-PG-001": "SAKURA-12 (Done, no commit/PR)",
        "ANOM-PG-002": "SAKURA-15 (In Progress, PR #5 merged)",
        "ANOM-PG-003": "P2-T05 / SAKURA-20 (planned end 2026-05-18, not Done)",
        "ANOM-PG-004": "Phase 2 Development (planned ~85% vs actual ~50%)",
        "ANOM-PG-005": "SAKURA-30 (orphan, no WBS parent)",
        "ANOM-BG-001": "SAKURA-25 (bug open since 2026-04-20)",
        "ANOM-BG-002": "SAKURA-26 (Severity-1 unresolved)",
        "ANOM-BG-003": "SAKURA-27/28/29 reopened",
        "ANOM-RK-001": "Slack C-004 blocker not in Jira",
        "ANOM-RK-002": "Guest-checkout decision not actioned",
        "ANOM-SC-001": "Sprint 7 & Sprint 8 both below 70%",
        "ANOM-QL-001": "PR #7 merged with no reviewer",
        "ANOM-QL-002": "PR #8 failing CI on main",
        "ANOM-RS-001": "Tanaka closed >40% of completed story points",
        "ANOM-CS-001": "Customer load-test question unanswered",
    },
}


def main() -> None:
    issues = build_issues()
    write_jira(issues)
    write_wbs()
    write_slack()
    write_github()
    write_minutes()
    (SAMPLE / "_ground_truth.json").write_text(
        json.dumps(GROUND_TRUTH, indent=2, ensure_ascii=False), encoding="utf-8")
    print(f"Sample data written to {SAMPLE}")
    print(f"  Jira issues : {len(issues)}")
    print(f"  Sprints     : {len(SPRINTS)}")
    print(f"  Seeded rules: {len(GROUND_TRUTH['expected_rules'])}")


if __name__ == "__main__":
    main()
