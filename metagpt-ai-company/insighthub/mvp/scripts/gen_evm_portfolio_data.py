"""Generate supplemental EVM-oriented sample data for DM portfolio UAT.

This dataset is separate from the default Sakura MVP sample. It models two
customer projects under one delivery manager:

- Project Atlas: large project, about 5M USD budget
- Project Beacon: medium project, about 1M USD budget

The files follow the same ingestion contracts as the core MVP:
- Jira export in XLSX
- WBS in XLSX
- Slack messages in JSON
- GitHub activity in JSON
- Meeting minutes in DOCX
- Previous reports in Markdown

Additionally, each project receives an `evm_context.yaml` file with human-
readable EVM metrics used for UAT review and PM interpretation.
"""

from __future__ import annotations

import json
from dataclasses import dataclass
from datetime import date
from pathlib import Path

import openpyxl
import yaml
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DATA_ROOT = ROOT / "data" / "evm-portfolio"


def d(value: str) -> date:
    return date.fromisoformat(value)


@dataclass(frozen=True)
class ProjectSpec:
    slug: str
    project_name: str
    budget_usd: int
    client_name: str
    report_start: str
    report_end: str
    jira_prefix: str
    repo_name: str
    channel: str
    status: str
    pv_usd: int
    ev_usd: int
    ac_usd: int
    eac_usd: int
    vac_usd: int
    bac_mm: float
    ev_mm: float
    ac_mm: float
    summary: str
    risk_summary: str
    milestone_summary: str


PROJECTS = [
    ProjectSpec(
        slug="project-atlas-5m",
        project_name="Project Atlas",
        budget_usd=5_000_000,
        client_name="Northwind Retail Japan",
        report_start="2026-06-01",
        report_end="2026-06-07",
        jira_prefix="ATLAS",
        repo_name="northwind/project-atlas",
        channel="#atlas-front-pm",
        status="Red",
        pv_usd=2_800_000,
        ev_usd=2_100_000,
        ac_usd=3_050_000,
        eac_usd=5_950_000,
        vac_usd=-950_000,
        bac_mm=125.0,
        ev_mm=52.5,
        ac_mm=61.0,
        summary="Large-scale commerce migration with CPI and SPI both below 1.0, plus baseline pressure from approved scope changes.",
        risk_summary="Critical path is slipping on payment cutover and warehouse sync, with cost overrun already visible.",
        milestone_summary="Wave 2 SIT exit is at risk; customer steering committee requested an EAC and mitigation update.",
    ),
    ProjectSpec(
        slug="project-beacon-1m",
        project_name="Project Beacon",
        budget_usd=1_000_000,
        client_name="Aster Finance Services",
        report_start="2026-06-01",
        report_end="2026-06-07",
        jira_prefix="BEACON",
        repo_name="aster/project-beacon",
        channel="#beacon-front-pm",
        status="Yellow",
        pv_usd=620_000,
        ev_usd=560_000,
        ac_usd=610_000,
        eac_usd=1_070_000,
        vac_usd=-70_000,
        bac_mm=28.0,
        ev_mm=15.5,
        ac_mm=17.0,
        summary="Mid-size portal refresh with mild CPI pressure but generally recoverable through scope discipline and tighter UAT control.",
        risk_summary="API certification and reporting dashboard quality remain the main near-term concerns.",
        milestone_summary="User acceptance preparation is progressing, but one approved change request has not yet been reflected in the baseline narrative.",
    ),
]


def main() -> None:
    DATA_ROOT.mkdir(parents=True, exist_ok=True)
    _write_root_readme()
    for spec in PROJECTS:
        _write_project(spec)
    print(f"Generated EVM portfolio samples under {DATA_ROOT}")


def _write_root_readme() -> None:
    content = """# EVM Portfolio Sample Pack

This supplemental sample pack is for UAT scenarios where a Delivery Manager
reviews multiple projects with earned-value-management signals.

Projects included:

- `project-atlas-5m`: large project, budget about 5M USD
- `project-beacon-1m`: medium project, budget about 1M USD

Artifacts per project:

- `connections.yaml`
- `evm_context.yaml`
- `Sample_Jira_Export.xlsx`
- `Sample_WBS.xlsx`
- `Sample_Slack_Messages.json`
- `Sample_GitHub_Activity.json`
- `minutes/*.docx`
- `previous_reports/*.md`

Suggested commands:

```powershell
& "C:\\Users\\HUY\\AppData\\Local\\Programs\\Python\\Python311\\python.exe" -m insighthub generate --type weekly --lang en --no-llm --connections data/evm-portfolio/project-atlas-5m/connections.yaml --out output-evm-atlas
& "C:\\Users\\HUY\\AppData\\Local\\Programs\\Python\\Python311\\python.exe" -m insighthub generate --type portfolio --lang en --no-llm --connections data/evm-portfolio/project-atlas-5m/connections.yaml data/evm-portfolio/project-beacon-1m/connections.yaml --out output-evm-portfolio
```
"""
    (DATA_ROOT / "README.md").write_text(content, encoding="utf-8")


def _write_project(spec: ProjectSpec) -> None:
    project_dir = DATA_ROOT / spec.slug
    minutes_dir = project_dir / "minutes"
    previous_dir = project_dir / "previous_reports"
    project_dir.mkdir(parents=True, exist_ok=True)
    minutes_dir.mkdir(parents=True, exist_ok=True)
    previous_dir.mkdir(parents=True, exist_ok=True)

    _write_connections(spec, project_dir)
    _write_evm_context(spec, project_dir)
    _write_jira(spec, project_dir)
    _write_wbs(spec, project_dir)
    _write_slack(spec, project_dir)
    _write_github(spec, project_dir)
    _write_minutes(spec, minutes_dir)
    _write_previous_reports(spec, previous_dir)


def _write_connections(spec: ProjectSpec, project_dir: Path) -> None:
    payload = {
        "project": spec.project_name,
        "period": {"start": spec.report_start, "end": spec.report_end},
        "base_urls": {
            "jira": "https://example.atlassian.net/browse",
            "github": f"https://github.com/{spec.repo_name}",
        },
        "systems": {
            "jira": {
                "adapter": "file",
                "auth": "env:JIRA_TOKEN",
                "file": f"data/evm-portfolio/{spec.slug}/Sample_Jira_Export.xlsx",
                "api_base": "https://example.atlassian.net",
            },
            "chat": {
                "adapter": "file",
                "auth": "env:SLACK_TOKEN",
                "provider": "slack",
                "file": f"data/evm-portfolio/{spec.slug}/Sample_Slack_Messages.json",
                "api_base": "https://slack.com/api",
            },
            "github": {
                "adapter": "file",
                "auth": "env:GITHUB_TOKEN",
                "file": f"data/evm-portfolio/{spec.slug}/Sample_GitHub_Activity.json",
                "api_base": "https://api.github.com",
            },
        },
        "uploads": {
            "wbs": f"data/evm-portfolio/{spec.slug}/Sample_WBS.xlsx",
            "minutes_dir": f"data/evm-portfolio/{spec.slug}/minutes",
        },
        "previous_reports_dir": f"data/evm-portfolio/{spec.slug}/previous_reports",
    }
    (project_dir / "connections.yaml").write_text(
        yaml.safe_dump(payload, sort_keys=False),
        encoding="utf-8",
    )


def _write_evm_context(spec: ProjectSpec, project_dir: Path) -> None:
    spi = round(spec.ev_usd / spec.pv_usd, 2)
    cpi = round(spec.ev_usd / spec.ac_usd, 2)
    payload = {
        "project": spec.project_name,
        "client_name": spec.client_name,
        "portfolio_role": "Delivery Manager",
        "budget_usd": spec.budget_usd,
        "report_period": {"start": spec.report_start, "end": spec.report_end},
        "evm": {
            "bac_usd": spec.budget_usd,
            "pv_usd": spec.pv_usd,
            "ev_usd": spec.ev_usd,
            "ac_usd": spec.ac_usd,
            "spi": spi,
            "cpi": cpi,
            "eac_usd": spec.eac_usd,
            "vac_usd": spec.vac_usd,
            "bac_mm": spec.bac_mm,
            "ev_mm": spec.ev_mm,
            "ac_mm": spec.ac_mm,
        },
        "dm_pm_kpis": _dm_pm_kpis(spec),
        "uat_focus": [
            "baseline control",
            "cost variance visibility",
            "schedule performance visibility",
            "change request impact",
            "milestone forecast credibility",
        ],
        "narrative_seed": {
            "summary": spec.summary,
            "risk_summary": spec.risk_summary,
            "milestone_summary": spec.milestone_summary,
        },
    }
    (project_dir / "evm_context.yaml").write_text(
        yaml.safe_dump(payload, sort_keys=False),
        encoding="utf-8",
    )


def _dm_pm_kpis(spec: ProjectSpec) -> dict[str, object]:
    spi = round(spec.ev_usd / spec.pv_usd, 2)
    cpi = round(spec.ev_usd / spec.ac_usd, 2)
    if spec.jira_prefix == "ATLAS":
        return {
            "schedule": {
                "planned_progress_pct": 56.0,
                "actual_progress_pct": 42.0,
                "schedule_variance_days": -14,
                "milestone_hit_rate_pct": 67.0,
                "sprint_commitment_pct": 55.0,
                "critical_path_tasks_at_risk": 2,
            },
            "cost": {
                "evm_is_primary": True,
                "spi": spi,
                "cpi": cpi,
                "cost_variance_usd": spec.ev_usd - spec.ac_usd,
                "schedule_variance_usd": spec.ev_usd - spec.pv_usd,
                "forecast_overrun_pct": 19.0,
                "burn_rate_pct_of_bac": 61.0,
            },
            "delivery": {
                "throughput_items": 3,
                "velocity_story_points": 16,
                "milestones_completed": 1,
                "milestones_due_this_month": 3,
            },
            "quality": {
                "open_bugs": 1,
                "sev1_open_bugs": 1,
                "bug_reopen_rate_pct": 18.0,
                "review_coverage_pct": 75.0,
                "ci_pass_rate_pct": 75.0,
            },
            "risk_blocker": {
                "open_blockers": 2,
                "max_blocker_age_days": 6,
                "unanswered_customer_questions": 2,
            },
            "change_control": {
                "approved_change_requests": 1,
                "pending_change_requests": 0,
                "change_budget_impact_usd": 350000,
                "change_schedule_impact_days": 10,
            },
            "resource": {
                "team_size": 18,
                "billable_utilization_pct": 88.0,
                "single_owner_dependency_pct": 31.0,
                "overtime_risk": "High",
            },
        }
    return {
        "schedule": {
            "planned_progress_pct": 82.0,
            "actual_progress_pct": 74.0,
            "schedule_variance_days": -4,
            "milestone_hit_rate_pct": 80.0,
            "sprint_commitment_pct": 60.0,
            "critical_path_tasks_at_risk": 1,
        },
        "cost": {
            "evm_is_primary": True,
            "spi": spi,
            "cpi": cpi,
            "cost_variance_usd": spec.ev_usd - spec.ac_usd,
            "schedule_variance_usd": spec.ev_usd - spec.pv_usd,
            "forecast_overrun_pct": 7.0,
            "burn_rate_pct_of_bac": 61.0,
        },
        "delivery": {
            "throughput_items": 3,
            "velocity_story_points": 10,
            "milestones_completed": 1,
            "milestones_due_this_month": 2,
        },
        "quality": {
            "open_bugs": 1,
            "sev1_open_bugs": 0,
            "bug_reopen_rate_pct": 8.0,
            "review_coverage_pct": 67.0,
            "ci_pass_rate_pct": 100.0,
        },
        "risk_blocker": {
            "open_blockers": 1,
            "max_blocker_age_days": 3,
            "unanswered_customer_questions": 1,
        },
        "change_control": {
            "approved_change_requests": 1,
            "pending_change_requests": 0,
            "change_budget_impact_usd": 45000,
            "change_schedule_impact_days": 3,
        },
        "resource": {
            "team_size": 7,
            "billable_utilization_pct": 84.0,
            "single_owner_dependency_pct": 26.0,
            "overtime_risk": "Medium",
        },
    }


def _write_jira(spec: ProjectSpec, project_dir: Path) -> None:
    prefix = spec.jira_prefix
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Issues"
    ws.append(
        [
            "Key",
            "Summary",
            "Type",
            "Status",
            "Priority",
            "Assignee",
            "Story Points",
            "Created",
            "Updated",
            "Resolved",
            "Sprint",
            "Labels",
            "Reopened",
        ]
    )

    rows = _jira_rows(spec)
    for row in rows:
        ws.append(row)

    sp = wb.create_sheet("Sprints")
    sp.append(["Sprint", "Start", "End", "Committed SP", "Completed SP"])
    for sprint in _sprints(spec):
        sp.append(sprint)

    wb.save(project_dir / "Sample_Jira_Export.xlsx")


def _jira_rows(spec: ProjectSpec) -> list[list[object]]:
    prefix = spec.jira_prefix
    if prefix == "ATLAS":
        return [
            [f"{prefix}-1", "Wave 2 payment cutover plan", "Story", "Done", "High", "Sato", 8, d("2026-05-20"), d("2026-06-02"), d("2026-06-02"), "Atlas Sprint 12", "", "No"],
            [f"{prefix}-2", "Warehouse sync refactor", "Story", "In Progress", "High", "Tanaka", 13, d("2026-05-18"), d("2026-06-07"), None, "Atlas Sprint 12", "critical-path", "No"],
            [f"{prefix}-3", "Cost dashboard EVM feed", "Task", "Done", "Medium", "Nguyen", 5, d("2026-05-22"), d("2026-06-03"), d("2026-06-03"), "Atlas Sprint 12", "evm", "No"],
            [f"{prefix}-4", "Approved Change Request 17 pricing update", "Story", "In Progress", "High", "Suzuki", 8, d("2026-05-25"), d("2026-06-06"), None, "Atlas Sprint 12", "change-request", "No"],
            [f"{prefix}-5", "SIT defect batch A", "Bug", "In Progress", "Severity-1", "Pham", 3, d("2026-05-28"), d("2026-06-07"), None, "Atlas Sprint 12", "sit", "No"],
            [f"{prefix}-6", "Legacy order migration rehearsal", "Task", "Done", "Medium", "Tanaka", 5, d("2026-05-19"), d("2026-06-01"), d("2026-06-01"), "Atlas Sprint 12", "migration", "No"],
            [f"{prefix}-7", "Carrier contract mapping", "Task", "In Progress", "Medium", "Nguyen", 5, d("2026-05-24"), d("2026-06-07"), None, "Atlas Sprint 12", "", "Yes"],
            [f"{prefix}-8", "Customer steering deck numbers", "Task", "Done", "Medium", "Sato", 3, d("2026-05-30"), d("2026-06-04"), d("2026-06-04"), "Atlas Sprint 12", "evm", "No"],
        ]
    return [
        [f"{prefix}-1", "Partner portal UAT checklist", "Task", "Done", "Medium", "Ito", 5, d("2026-05-20"), d("2026-06-02"), d("2026-06-02"), "Beacon Sprint 5", "", "No"],
        [f"{prefix}-2", "API certification package", "Story", "In Progress", "High", "Yamada", 8, d("2026-05-24"), d("2026-06-07"), None, "Beacon Sprint 5", "external-dependency", "No"],
        [f"{prefix}-3", "Approved Change Request 04 reporting tweak", "Task", "In Progress", "Medium", "Nguyen", 3, d("2026-05-29"), d("2026-06-06"), None, "Beacon Sprint 5", "change-request", "No"],
        [f"{prefix}-4", "Dashboard widget styling", "Task", "Done", "Low", "Ito", 3, d("2026-05-26"), d("2026-06-03"), d("2026-06-03"), "Beacon Sprint 5", "", "No"],
        [f"{prefix}-5", "Regression defect in PDF export", "Bug", "In Progress", "High", "Sato", 2, d("2026-05-31"), d("2026-06-07"), None, "Beacon Sprint 5", "uat", "No"],
        [f"{prefix}-6", "Budget narrative alignment", "Task", "Done", "Medium", "Nguyen", 2, d("2026-05-28"), d("2026-06-04"), d("2026-06-04"), "Beacon Sprint 5", "evm", "No"],
    ]


def _sprints(spec: ProjectSpec) -> list[list[object]]:
    if spec.jira_prefix == "ATLAS":
        return [
            ["Atlas Sprint 10", d("2026-04-27"), d("2026-05-10"), 48, 37],
            ["Atlas Sprint 11", d("2026-05-11"), d("2026-05-24"), 52, 34],
            ["Atlas Sprint 12", d("2026-05-25"), d("2026-06-07"), 56, 31],
        ]
    return [
        ["Beacon Sprint 3", d("2026-04-27"), d("2026-05-10"), 18, 16],
        ["Beacon Sprint 4", d("2026-05-11"), d("2026-05-24"), 20, 14],
        ["Beacon Sprint 5", d("2026-05-25"), d("2026-06-07"), 22, 12],
    ]


def _write_wbs(spec: ProjectSpec, project_dir: Path) -> None:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "WBS"
    ws.append(["Task ID", "Phase", "Task Name", "Planned Start", "Planned End", "Planned MM", "Planned %", "Owner", "Jira Key"])
    for row in _wbs_rows(spec):
        ws.append(row)
    wb.save(project_dir / "Sample_WBS.xlsx")


def _wbs_rows(spec: ProjectSpec) -> list[list[object]]:
    prefix = spec.jira_prefix
    if prefix == "ATLAS":
        return [
            ["A1-T01", "Wave 1", "Migration readiness", d("2026-04-01"), d("2026-05-15"), 20.0, 100, "Sato", ""],
            ["A2-T01", "Wave 2 SIT", "Payment cutover plan", d("2026-05-18"), d("2026-06-03"), 18.0, 95, "Sato", f"{prefix}-1"],
            ["A2-T02", "Wave 2 SIT", "Warehouse sync refactor", d("2026-05-20"), d("2026-06-04"), 22.0, 90, "Tanaka", f"{prefix}-2"],
            ["A2-T03", "Wave 2 SIT", "Cost dashboard EVM feed", d("2026-05-22"), d("2026-06-02"), 10.0, 100, "Nguyen", f"{prefix}-3"],
            ["A2-T04", "Wave 2 SIT", "Pricing CR implementation", d("2026-05-25"), d("2026-06-05"), 16.0, 75, "Suzuki", f"{prefix}-4"],
            ["A2-T05", "Wave 2 SIT", "Migration rehearsal", d("2026-05-19"), d("2026-06-01"), 12.0, 100, "Tanaka", f"{prefix}-6"],
            ["A3-T01", "UAT", "Steering deck numbers", d("2026-05-30"), d("2026-06-04"), 6.0, 100, "Sato", f"{prefix}-8"],
            ["A3-T02", "UAT", "Carrier contract mapping", d("2026-05-24"), d("2026-06-02"), 8.0, 85, "Nguyen", f"{prefix}-7"],
        ]
    return [
        ["B1-T01", "Foundation", "Partner portal baseline", d("2026-04-20"), d("2026-05-20"), 5.0, 100, "Ito", ""],
        ["B2-T01", "Certification", "API certification package", d("2026-05-24"), d("2026-06-04"), 8.0, 90, "Yamada", f"{prefix}-2"],
        ["B2-T02", "Certification", "Reporting CR update", d("2026-05-29"), d("2026-06-03"), 3.0, 100, "Nguyen", f"{prefix}-3"],
        ["B3-T01", "UAT", "Partner portal UAT checklist", d("2026-05-20"), d("2026-06-02"), 4.0, 100, "Ito", f"{prefix}-1"],
        ["B3-T02", "UAT", "PDF export regression fix", d("2026-05-31"), d("2026-06-05"), 2.5, 80, "Sato", f"{prefix}-5"],
        ["B3-T03", "UAT", "Budget narrative alignment", d("2026-05-28"), d("2026-06-04"), 1.5, 100, "Nguyen", f"{prefix}-6"],
    ]


def _write_slack(spec: ProjectSpec, project_dir: Path) -> None:
    messages = _slack_messages(spec)
    payload = {
        "channel": spec.channel,
        "messages": [
            {"id": item["id"], "channel": spec.channel, "user": item["user"], "ts": item["ts"], "text": item["text"]}
            for item in messages
        ],
    }
    (project_dir / "Sample_Slack_Messages.json").write_text(
        json.dumps(payload, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )


def _slack_messages(spec: ProjectSpec) -> list[dict[str, str]]:
    prefix = spec.jira_prefix
    if prefix == "ATLAS":
        return [
            {"id": "AT-C001", "user": "Tanaka", "ts": "2026-06-02T09:10:00", "text": f"{prefix}-2 warehouse sync is still blocking the SIT exit path and AC is already above the last forecast."},
            {"id": "AT-C002", "user": "Sato", "ts": "2026-06-03T10:00:00", "text": f"{prefix}-1 cutover deck is ready but the customer asked for CPI and SPI in the steering committee narrative."},
            {"id": "AT-C003", "user": "Nguyen", "ts": "2026-06-04T13:20:00", "text": f"{prefix}-3 completed. EVM feed now matches the reporting model."},
            {"id": "AT-C004", "user": "Pham", "ts": "2026-06-05T14:45:00", "text": "Blocker: customer sandbox batch file is delayed again, which could push the payment cutover milestone."},
            {"id": "AT-C005", "user": "Suzuki", "ts": "2026-06-06T11:05:00", "text": f"{prefix}-5 severity-1 SIT defect remains open and mitigation is not yet accepted by the customer."},
        ]
    return [
        {"id": "BC-C001", "user": "Ito", "ts": "2026-06-02T09:00:00", "text": f"{prefix}-1 UAT checklist completed and shared with the customer."},
        {"id": "BC-C002", "user": "Yamada", "ts": "2026-06-03T15:10:00", "text": f"{prefix}-2 certification package still depends on the vendor approval window."},
            {"id": "BC-C003", "user": "Nguyen", "ts": "2026-06-05T16:00:00", "text": "Risk: approved Change Request 04 changed the dashboard wording, but the customer-facing budget note still uses the old baseline."},
    ]


def _write_github(spec: ProjectSpec, project_dir: Path) -> None:
    payload = _github_payload(spec)
    (project_dir / "Sample_GitHub_Activity.json").write_text(
        json.dumps(payload, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )


def _github_payload(spec: ProjectSpec) -> dict[str, object]:
    prefix = spec.jira_prefix
    if prefix == "ATLAS":
        commits = [
            {"sha": "aa11bb1", "message": f"{prefix}-1 finalize cutover task breakdown", "author": "Sato", "date": "2026-06-02T08:10:00"},
            {"sha": "aa11bb2", "message": f"{prefix}-3 wire EVM fact feed into report exporter", "author": "Nguyen", "date": "2026-06-03T08:45:00"},
            {"sha": "aa11bb3", "message": f"{prefix}-6 migration rehearsal scripts", "author": "Tanaka", "date": "2026-06-01T17:20:00"},
            {"sha": "aa11bb4", "message": f"{prefix}-8 steering committee finance appendix", "author": "Sato", "date": "2026-06-04T18:00:00"},
        ]
        prs = [
            {"number": 21, "title": f"{prefix}-1 payment cutover plan", "author": "Sato", "reviewers": ["Tanaka"], "state": "merged", "created_at": "2026-06-01T09:00:00", "merged_at": "2026-06-02T12:00:00", "ci_status": "success"},
            {"number": 22, "title": f"{prefix}-3 EVM feed wiring", "author": "Nguyen", "reviewers": ["Suzuki"], "state": "merged", "created_at": "2026-06-02T13:00:00", "merged_at": "2026-06-03T16:00:00", "ci_status": "success"},
            {"number": 23, "title": f"{prefix}-2 warehouse sync refactor", "author": "Tanaka", "reviewers": ["Sato"], "state": "open", "created_at": "2026-06-03T10:00:00", "merged_at": None, "ci_status": "failing"},
            {"number": 24, "title": f"{prefix}-4 approved pricing change request", "author": "Suzuki", "reviewers": ["Nguyen"], "state": "merged", "created_at": "2026-06-04T11:00:00", "merged_at": "2026-06-06T14:00:00", "ci_status": "success"},
        ]
        return {"repo": spec.repo_name, "default_branch": "main", "commits": commits, "pull_requests": prs}
    commits = [
        {"sha": "cc22dd1", "message": f"{prefix}-1 finalize UAT checklist", "author": "Ito", "date": "2026-06-02T08:15:00"},
        {"sha": "cc22dd2", "message": f"{prefix}-4 dashboard widget styling", "author": "Ito", "date": "2026-06-03T09:20:00"},
        {"sha": "cc22dd3", "message": f"{prefix}-6 align budget narrative wording", "author": "Nguyen", "date": "2026-06-04T17:00:00"},
    ]
    prs = [
        {"number": 11, "title": f"{prefix}-1 partner portal UAT checklist", "author": "Ito", "reviewers": ["Nguyen"], "state": "merged", "created_at": "2026-06-01T09:00:00", "merged_at": "2026-06-02T12:30:00", "ci_status": "success"},
        {"number": 12, "title": f"{prefix}-2 certification package", "author": "Yamada", "reviewers": ["Sato"], "state": "open", "created_at": "2026-06-03T10:00:00", "merged_at": None, "ci_status": "success"},
        {"number": 13, "title": f"{prefix}-3 approved reporting change request", "author": "Nguyen", "reviewers": [], "state": "merged", "created_at": "2026-06-04T11:00:00", "merged_at": "2026-06-05T14:00:00", "ci_status": "success"},
    ]
    return {"repo": spec.repo_name, "default_branch": "main", "commits": commits, "pull_requests": prs}


def _write_minutes(spec: ProjectSpec, minutes_dir: Path) -> None:
    minutes = _minutes_payload(spec)
    for filename, payload in minutes:
        document = Document()
        document.add_paragraph(f"Date: {payload['date']}")
        document.add_heading("Decisions", level=1)
        for line in payload["decisions"]:
            document.add_paragraph(line, style="List Bullet")
        document.add_heading("Action Items", level=1)
        for line in payload["actions"]:
            document.add_paragraph(line, style="List Bullet")
        document.add_heading("Questions", level=1)
        for line in payload["questions"]:
            document.add_paragraph(line, style="List Bullet")
        document.save(minutes_dir / filename)


def _minutes_payload(spec: ProjectSpec) -> list[tuple[str, dict[str, object]]]:
    prefix = spec.jira_prefix
    if prefix == "ATLAS":
        return [
            (
                "minutes_2026-06-03_steering-committee.docx",
                {
                    "date": "2026-06-03",
                    "decisions": [
                        "Approved Change Request 17 must be included in the revised cost baseline before the next steering review.",
                    ],
                    "actions": [
                        "Tanaka to prepare an updated EAC and recovery plan by 2026-06-08.",
                    ],
                    "questions": [
                        "How much budget contingency remains if Wave 2 SIT exit slips by two weeks?",
                    ],
                },
            ),
            (
                "minutes_2026-06-05_customer-sync.docx",
                {
                    "date": "2026-06-05",
                    "decisions": [
                        "Customer asked that CPI and SPI be explicitly called out in the weekly steering summary.",
                    ],
                    "actions": [
                        "Nguyen to reconcile dashboard numbers with the approved baseline by 2026-06-09.",
                    ],
                    "questions": [
                        "Can the payment cutover milestone still be met without overtime spending?",
                    ],
                },
            ),
        ]
    return [
        (
            "minutes_2026-06-03_delivery-sync.docx",
            {
                "date": "2026-06-03",
                "decisions": [
                    "Change Request 04 is approved but should not change the committed UAT date unless certification slips.",
                ],
                "actions": [
                    "Yamada to confirm vendor certification timing by 2026-06-06.",
                ],
                "questions": [
                    "Should the customer-facing budget narrative mention the updated EAC this week or next week?",
                ],
            },
        ),
    ]


def _write_previous_reports(spec: ProjectSpec, previous_dir: Path) -> None:
    spi = round(spec.ev_usd / spec.pv_usd, 2)
    cpi = round(spec.ev_usd / spec.ac_usd, 2)
    monthly = [
        (
            "monthly_2026-05.md",
            f"""# {spec.project_name} - Monthly Status Report

Period: 2026-05-01 - 2026-05-31
Overall status: {spec.status}

## Executive Summary

- Budget baseline: {spec.budget_usd:,} USD.
- May close showed SPI {spi:.2f} and CPI {cpi:.2f}.
- Main PM message: {spec.summary}
""",
        ),
        (
            "monthly_2026-04.md",
            f"""# {spec.project_name} - Monthly Status Report

Period: 2026-04-01 - 2026-04-30
Overall status: Yellow

## Executive Summary

- Baseline was still broadly intact, but early warning signs were visible in cost and milestone confidence.
- Customer requested tighter narrative around milestone risk and change-control impact.
""",
        ),
        (
            "monthly_2026-03.md",
            f"""# {spec.project_name} - Monthly Status Report

Period: 2026-03-01 - 2026-03-31
Overall status: Green

## Executive Summary

- Project mobilization and baseline planning completed.
- No material cost variance was visible yet.
""",
        ),
    ]
    weekly = [
        (
            "weekly_2026-05-31.md",
            f"""# {spec.project_name} - Weekly Status Report

Period: 2026-05-25 - 2026-05-31
Overall status: {spec.status}

## Executive Summary

- Prior week focus: {spec.milestone_summary}
- PM note: recovery action depends on baseline clarity and owner accountability.
""",
        ),
        (
            "weekly_2026-05-24.md",
            f"""# {spec.project_name} - Weekly Status Report

Period: 2026-05-18 - 2026-05-24
Overall status: Yellow

## Executive Summary

- The project entered the current reporting cycle with visible schedule pressure.
- Cost trend was manageable but already moving away from the original forecast.
""",
        ),
        (
            "weekly_2026-05-17.md",
            f"""# {spec.project_name} - Weekly Status Report

Period: 2026-05-11 - 2026-05-17
Overall status: Yellow

## Executive Summary

- Change requests began to affect near-term planning assumptions.
- EVM interpretation was requested for the next steering review.
""",
        ),
        (
            "weekly_2026-05-10.md",
            f"""# {spec.project_name} - Weekly Status Report

Period: 2026-05-04 - 2026-05-10
Overall status: Green

## Executive Summary

- Delivery baseline still appeared achievable.
- Main concern was ensuring execution data and PM narrative stayed aligned.
""",
        ),
    ]
    for name, content in [*monthly, *weekly]:
        (previous_dir / name).write_text(content, encoding="utf-8")


if __name__ == "__main__":
    main()
