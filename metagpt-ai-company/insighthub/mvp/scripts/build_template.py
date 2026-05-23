"""Build the weekly DOCX report template."""

from __future__ import annotations

from pathlib import Path

from docx import Document


SECTIONS: tuple[tuple[str, str], ...] = (
    ("exec_summary", "1. Executive Summary"),
    ("progress", "2. Progress"),
    ("completed", "3. Completed"),
    ("in_progress", "4. In Progress"),
    ("next_week", "5. Next Week"),
    ("blockers", "6. Blockers"),
    ("bugs", "7. Bugs"),
    ("decisions", "8. Decisions"),
    ("metrics", "9. Metrics"),
)


def build_template(path: str | Path = "templates/weekly_template.docx") -> Path:
    """Create the weekly report template and return its path."""

    template_path = Path(path)
    template_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    document.add_heading("{{project_name}} — Weekly Status Report", level=0)
    document.add_paragraph("Period: {{period}}")
    document.add_paragraph("Overall Status: {{overall_status}}")

    for section_id, title in SECTIONS:
        document.add_heading(title, level=1)
        document.add_paragraph(f"{{{{{section_id}}}}}")

    document.save(template_path)
    return template_path


if __name__ == "__main__":
    build_template()
