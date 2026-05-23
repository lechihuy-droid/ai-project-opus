"""Export validated reports to DOCX, Markdown, traceability, and audit files."""

from __future__ import annotations

import json
import logging
import os
import re
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from insighthub import i18n
from insighthub.schema import Fact, Facts, Report
from insighthub.templating import SECTIONS, render, resolve_template

logger = logging.getLogger(__name__)
CITATION_RE = re.compile(r"\[([a-zA-Z0-9_-]+):([^\]]+)\]")


def export(
    report: Report,
    facts: Facts,
    out_dir: str = "output",
    template_name: str | None = None,
    diff_markdown: str | None = None,
) -> dict[str, str]:
    """Export report artifacts and return their file paths."""

    output_dir = Path(out_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    prefix = _output_prefix(report)
    docx_path = output_dir / f"{prefix}.docx"
    md_path = output_dir / f"{prefix}.md"
    traceability_path = output_dir / "traceability.json"
    audit_path = output_dir / "audit_log.md"
    diff_path = output_dir / "diff.md"

    citation_urls = _citation_urls(facts)
    document = _render_document(report, template_name)
    _localize_docx_header(document, report)
    _link_docx_citations(document, citation_urls)
    _link_docx_markdown_links(document)
    document.save(docx_path)
    pdf_path = _to_pdf(str(docx_path))

    md_path.write_text(_build_markdown(report, citation_urls), encoding="utf-8")
    traceability_path.write_text(
        json.dumps(_build_traceability(facts), ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    audit_text = _build_audit_log(report, facts)
    if audit_path.exists():
        existing_audit = audit_path.read_text(encoding="utf-8").rstrip()
        audit_path.write_text(f"{existing_audit}\n\n{audit_text}", encoding="utf-8")
    else:
        audit_path.write_text(audit_text, encoding="utf-8")

    if diff_markdown is not None:
        diff_path.write_text(_link_markdown_citations(diff_markdown, citation_urls), encoding="utf-8")

    paths = {
        f"{prefix}_docx": str(docx_path),
        f"{prefix}_md": str(md_path),
        "traceability_json": str(traceability_path),
        "audit_log_md": str(audit_path),
    }
    if diff_markdown is not None:
        paths["diff_md"] = str(diff_path)
    if pdf_path:
        paths[f"{prefix}_pdf"] = pdf_path
    return paths


def _ensure_template_exists(template_path: Path) -> None:
    if template_path.exists():
        return
    if template_path == Path("templates/weekly_template.docx"):
        from scripts.build_template import build_template

        build_template(template_path)
        return
    raise FileNotFoundError(f"Template not found: {template_path}")


def _render_document(report: Report, template_name: str | None):
    if report.report_type == "portfolio":
        return _render_portfolio_document(report)
    template = resolve_template(report.report_type, template_name)
    template_path = template.path
    _ensure_template_exists(template_path)
    return render(report, str(template_path))


def _render_portfolio_document(report: Report):
    document = Document()
    document.add_heading(f"{report.project_name} - Portfolio Dashboard", level=0)
    document.add_paragraph(f"Period: {report.period_start.isoformat()} - {report.period_end.isoformat()}")
    document.add_paragraph(f"Overall Status: {report.overall_status}")
    for section in report.sections:
        document.add_heading(section.title, level=1)
        for line in [item for item in section.body.splitlines() if item.strip()]:
            paragraph = document.add_paragraph()
            if line.lstrip().startswith("- "):
                paragraph.style = "List Bullet"
                paragraph.add_run(line.lstrip()[2:].strip())
            else:
                paragraph.add_run(line.strip())
    return document


def _to_pdf(docx_path: str) -> str | None:
    """Convert DOCX to PDF when docx2pdf/Word is available."""

    if os.environ.get("PYTEST_CURRENT_TEST"):
        return None

    pdf_path = str(Path(docx_path).with_suffix(".pdf"))
    try:
        from docx2pdf import convert

        convert(docx_path, pdf_path)
    except Exception as exc:
        logger.warning("PDF skipped: %s", exc)
        return None
    return pdf_path


def _build_markdown(report: Report, citation_urls: dict[str, str]) -> str:
    if report.report_type == "portfolio":
        return _build_portfolio_markdown(report, citation_urls)

    section_by_id = {section.section_id: section for section in report.sections}
    lines = [
        f"# {report.project_name} — {i18n.header('report_title', report.language)}",
        "",
        f"{i18n.header('period', report.language)}: {report.period_start.isoformat()} - {report.period_end.isoformat()}",
        f"{i18n.header('overall_status', report.language)}: {report.overall_status}",
        "",
    ]

    for section_id, default_title in SECTIONS:
        section = section_by_id.get(section_id)
        title = section.title if section else default_title
        body = section.body if section and section.body else "(no data)"
        lines.extend([f"## {title}", "", _link_markdown_citations(body, citation_urls), ""])

    weekly_section_ids = {section_id for section_id, _title in SECTIONS}
    for section in report.sections:
        if section.section_id in weekly_section_ids:
            continue
        title = section.title or section.section_id
        body = section.body or "(no data)"
        lines.extend([f"## {title}", "", _link_markdown_citations(body, citation_urls), ""])

    if report.report_type == "monthly":
        lines[0] = f"# {report.project_name} - Monthly Status Report"

    return "\n".join(lines).rstrip() + "\n"


def _build_portfolio_markdown(report: Report, citation_urls: dict[str, str]) -> str:
    lines = [
        f"# {report.project_name} - Portfolio Dashboard",
        "",
        f"Period: {report.period_start.isoformat()} - {report.period_end.isoformat()}",
        f"Overall Status: {report.overall_status}",
        "",
    ]
    for section in report.sections:
        title = section.title or section.section_id
        body = section.body or "(no data)"
        lines.extend([f"## {title}", "", _link_markdown_citations(body, citation_urls), ""])
    return "\n".join(lines).rstrip() + "\n"


def _localize_docx_header(document, report: Report) -> None:
    replacements = {
        "Weekly Status Report": _localized_report_title(report),
        "Period:": f"{i18n.header('period', report.language)}:",
        "Overall Status:": f"{i18n.header('overall_status', report.language)}:",
    }
    for paragraph in document.paragraphs:
        text = paragraph.text
        rendered = text
        for old, new in replacements.items():
            rendered = rendered.replace(old, new)
        if rendered != text:
            _set_paragraph_text(paragraph, rendered)


def _citation_urls(facts: Facts) -> dict[str, str]:
    urls: dict[str, str] = {}
    for fact in _iter_facts(facts):
        for citation in fact.citations:
            if citation.url:
                urls[citation.cite()] = citation.url
    return urls


def _link_markdown_citations(text: str, citation_urls: dict[str, str]) -> str:
    def replace(match: re.Match[str]) -> str:
        token = match.group(0)
        url = citation_urls.get(token)
        if not url:
            return token
        return f"[{match.group(1)}:{match.group(2)}]({url})"

    return CITATION_RE.sub(replace, text)


def _link_docx_citations(document, citation_urls: dict[str, str]) -> None:
    if not citation_urls:
        return
    for paragraph in document.paragraphs:
        if not CITATION_RE.search(paragraph.text):
            continue
        style = paragraph.style
        text = paragraph.text
        for run in paragraph.runs:
            run.text = ""
        cursor = 0
        for match in CITATION_RE.finditer(text):
            if match.start() > cursor:
                paragraph.add_run(text[cursor:match.start()])
            token = match.group(0)
            url = citation_urls.get(token)
            if url:
                _add_hyperlink(paragraph, token, url)
            else:
                paragraph.add_run(token)
            cursor = match.end()
        if cursor < len(text):
            paragraph.add_run(text[cursor:])
        paragraph.style = style


def _link_docx_markdown_links(document) -> None:
    markdown_link_re = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
    for paragraph in document.paragraphs:
        text = paragraph.text
        if not markdown_link_re.search(text):
            continue
        style = paragraph.style
        for run in paragraph.runs:
            run.text = ""
        cursor = 0
        for match in markdown_link_re.finditer(text):
            if match.start() > cursor:
                paragraph.add_run(text[cursor:match.start()])
            _add_hyperlink(paragraph, match.group(1), match.group(2))
            cursor = match.end()
        if cursor < len(text):
            paragraph.add_run(text[cursor:])
        paragraph.style = style


def _add_hyperlink(paragraph: Paragraph, text: str, url: str) -> None:
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(color)
    properties.append(underline)
    run.append(properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def _build_traceability(facts: Facts) -> dict[str, dict[str, object]]:
    traceability: dict[str, dict[str, object]] = {}
    for fact in _iter_facts(facts):
        traceability[fact.id] = {
            "label": fact.label,
            "value": fact.value,
            "citations": [
                {
                    "system": citation.system,
                    "ref_id": citation.ref_id,
                    "label": citation.label,
                    "url": citation.url,
                }
                for citation in fact.citations
            ],
        }
    return traceability


def _iter_facts(facts: Facts) -> list[Fact]:
    result: list[Fact] = []
    for section in facts.sections:
        result.extend(section.facts)
        result.extend(section.bullet_items)
    return result


def _build_audit_log(report: Report, facts: Facts) -> str:
    fact_count = len(_iter_facts(facts))
    section_ids = [section.section_id for section in report.sections]
    lines = [
        "# Audit Log",
        "",
        f"- Timestamp: {datetime.now(timezone.utc).isoformat()}",
        f"- Project: {report.project_name}",
        f"- Period: {report.period_start.isoformat()} - {report.period_end.isoformat()}",
        f"- Fact count: {fact_count}",
        f"- Anomaly count: {len(facts.anomalies)}",
        f"- Sections exported: {', '.join(section_ids)}",
        "- Model used: claude-opus-4-7",
        "",
    ]
    return "\n".join(lines)


def _localized_report_title(report: Report) -> str:
    if report.report_type == "portfolio":
        return "Portfolio Dashboard"
    if report.report_type == "monthly":
        return "Monthly Status Report"
    return i18n.header("report_title", report.language)


def _output_prefix(report: Report) -> str:
    if report.report_type == "monthly":
        return "monthly"
    if report.report_type == "portfolio":
        return "portfolio"
    return "weekly"
